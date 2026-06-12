"""
14_report.py
------------
Generate an IIRS-style project report (.docx) with python-docx.

Layout (A4, Times New Roman 12 pt, 1.5 line spacing, justified):
  Cover page
  Declaration / certificate placeholder
  Acknowledgement placeholder
  Abstract
  Table of contents
  List of figures
  List of tables
  1  Introduction
  2  Study Area
  3  Data and Software Used
  4  Methodology
  5  Results and Discussion
  6  Conclusion and Future Scope
  7  References  (APA 7)

The report wording is original prose authored for this project; all factual
statements about classifier accuracy are taken from the project artefacts
(D:/Dehradun_OBIA/06_validation/accuracy_summary.csv and area tables).
External claims are attributed to peer-reviewed sources (APA 7).
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

MAPS = Path(r"D:\Dehradun_OBIA\08_maps_final")
VAL = Path(r"D:\Dehradun_OBIA\06_validation")
CHG = Path(r"D:\Dehradun_OBIA\07_change_detection")
OBIA = Path(r"D:\Dehradun_OBIA\04_classification_OBIA")
OUT = Path(r"D:\Dehradun_OBIA\09_report_slides")
OUT.mkdir(parents=True, exist_ok=True)

REPORT_PATH = OUT / "Dehradun_OBIA_Report.docx"


# ----------------------------- low-level helpers -----------------------------

def set_default_style(doc: Document) -> None:
    """Times New Roman 12 pt, 1.5 spacing, justified — applied to Normal style."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_heading_style(doc: Document, level: int, size_pt: int, bold: bool = True) -> None:
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False,
             align="justify", size_pt: int = 12, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    h.paragraph_format.line_spacing = 1.5
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_image(doc, path: Path, width_in: float = 6.0, caption: str | None = None,
              fig_num: int | None = None) -> None:
    if not path.exists():
        add_para(doc, f"[Figure missing: {path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(f"Figure {fig_num}. {caption}" if fig_num else caption)
        crun.italic = True
        crun.font.name = "Times New Roman"
        crun.font.size = Pt(11)


def add_table_caption(doc, table_num: int, text: str) -> None:
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cp.add_run(f"Table {table_num}. {text}")
    crun.italic = True
    crun.font.name = "Times New Roman"
    crun.font.size = Pt(11)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = table.rows[i].cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def add_page_break(doc) -> None:
    doc.add_page_break()


def insert_page_field(paragraph) -> None:
    """Insert a Word PAGE field into a paragraph."""
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_a4_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    # Footer with page number
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_page_field(p)


# --------------------------------- content -----------------------------------

def cover_page(doc: Document) -> None:
    # vertical spacing
    add_para(doc, "", space_after=24)
    add_para(doc, "INDIAN INSTITUTE OF REMOTE SENSING", bold=True, align="center",
             size_pt=18, space_after=4)
    add_para(doc, "Indian Space Research Organisation (ISRO)", align="center",
             size_pt=12, space_after=4)
    add_para(doc, "Department of Space, Government of India", align="center",
             size_pt=12, space_after=4)
    add_para(doc, "Dehradun – 248001, Uttarakhand, India", align="center",
             italic=True, size_pt=11, space_after=60)

    add_para(doc, "MINI-PROJECT REPORT", bold=True, align="center",
             size_pt=14, space_after=20)
    add_para(doc,
             "Object-Based vs Pixel-Based Classification for "
             "Dehradun Urban Expansion Mapping (2016–2024)",
             bold=True, align="center", size_pt=18, space_after=40)

    add_para(doc, "A Comparative Assessment Using Sentinel-2 Surface Reflectance",
             italic=True, align="center", size_pt=13, space_after=80)

    add_para(doc, "Submitted by", italic=True, align="center", size_pt=12, space_after=6)
    add_para(doc, "Mohammad Salman", bold=True, align="center",
             size_pt=14, space_after=4)
    add_para(doc, "M.Sc. Geo-informatics", align="center", size_pt=12, space_after=2)
    add_para(doc, "Indian Institute of Remote Sensing, ISRO, Dehradun", align="center",
             italic=True, size_pt=11, space_after=60)

    add_para(doc, "Under the guidance of", italic=True, align="center",
             size_pt=12, space_after=4)
    add_para(doc, "[Supervisor Name]", bold=True, align="center",
             size_pt=13, space_after=2)
    add_para(doc, "[Department / Division]", italic=True, align="center",
             size_pt=11, space_after=80)

    add_para(doc, "June 2026", bold=True, align="center", size_pt=13)
    add_page_break(doc)


def declaration(doc: Document) -> None:
    add_heading(doc, "Declaration", 1)
    add_para(doc,
             "I, Mohammad Salman, hereby declare that the mini-project report titled "
             "“Object-Based vs Pixel-Based Classification for Dehradun Urban Expansion "
             "Mapping (2016–2024)” submitted to the Indian Institute of Remote Sensing, "
             "Dehradun, is a record of original work carried out by me under the guidance "
             "of my supervisor. The analyses, results, and interpretations presented in "
             "this report are my own, and all external sources of information, theory, "
             "and tools have been duly cited in the references using APA 7th edition "
             "style. No part of this report has been copied verbatim from any published "
             "source, and the work has not been submitted elsewhere for any degree or "
             "diploma.")
    add_para(doc, "", space_after=24)
    add_para(doc, "Place: Dehradun", align="left")
    add_para(doc, "Date: June 2026", align="left", space_after=24)
    add_para(doc, "Mohammad Salman", bold=True, align="right")
    add_para(doc, "M.Sc. Geo-informatics, IIRS", align="right", italic=True)
    add_page_break(doc)


def acknowledgement(doc: Document) -> None:
    add_heading(doc, "Acknowledgement", 1)
    add_para(doc,
             "I would like to express my sincere gratitude to my supervisor at the "
             "Indian Institute of Remote Sensing for the continuous guidance, "
             "constructive feedback, and encouragement received throughout this "
             "mini-project. I am thankful to the faculty of the Geo-informatics "
             "Department for the foundational training that made this comparative "
             "study possible. I also acknowledge the European Space Agency’s Copernicus "
             "programme for providing free access to Sentinel-2 surface-reflectance "
             "data, and the Google Earth Engine team for the platform that enabled "
             "cloud-free composite generation. Finally, I thank my classmates and "
             "family for their support during the project period.")
    add_para(doc, "", space_after=12)
    add_para(doc, "Mohammad Salman", bold=True, align="right")
    add_page_break(doc)


def abstract(doc: Document) -> None:
    add_heading(doc, "Abstract", 1)
    add_para(doc,
             "Dehradun, the capital of Uttarakhand, has experienced rapid peri-urban "
             "expansion over the past decade. Accurately mapping this expansion is a "
             "pre-requisite for evidence-based urban planning, hydrological assessment, "
             "and environmental impact studies. The present mini-project compares two "
             "supervised land-cover classification paradigms — pixel-based Random "
             "Forest and Object-Based Image Analysis (OBIA) with Random Forest on "
             "Simple Linear Iterative Clustering (SLIC) superpixels — applied to two "
             "co-registered Sentinel-2 surface-reflectance composites for the winter "
             "windows of November 2016 to February 2017 and November 2024 to February "
             "2025. A five-class land-cover scheme (built-up, dense vegetation, mixed "
             "vegetation/cropland, bare/open land, and water) is adopted, and identical "
             "training samples, spectral features (B2, B3, B4, B8, NDVI, NDBI) and "
             "Random Forest hyperparameters are used for both methods, isolating the "
             "spatial unit of classification as the only independent variable. Accuracy "
             "is assessed against 75 stratified random reference points per date "
             "interpreted in Google Earth Pro. Pixel-based classification achieved an "
             "overall accuracy of 78.7 % and a Cohen’s κ of 0.72 on both dates, whereas "
             "OBIA achieved 65.3 % (κ = 0.57) in 2016 and 69.3 % (κ = 0.62) in 2024. "
             "Both methods agree that built-up area expanded substantially over the "
             "study period; the pixel-based estimate of approximately 80 km² of new "
             "built-up land is preferred as the more accurate figure. The lower OBIA "
             "accuracy is attributed to the relatively coarse SLIC segments (median "
             "≈ 12.4 ha) averaging away within-segment spectral diversity at the "
             "urban–cropland fringe, where most of Dehradun’s actual change is "
             "occurring. The study concludes that, at the 10 m Sentinel-2 resolution "
             "and at a single segmentation scale, the visual smoothness of OBIA does "
             "not translate into higher accuracy for medium-resolution urban "
             "applications.")
    add_para(doc, "", space_after=6)
    add_para(doc,
             "Keywords: Sentinel-2; OBIA; SLIC; Random Forest; urban expansion; "
             "Dehradun; change detection; APA 7.", italic=True)
    add_page_break(doc)


def table_of_contents(doc: Document) -> None:
    add_heading(doc, "Table of Contents", 1)
    # Static TOC text — easy to read and verify offline.
    entries = [
        ("Declaration", "ii"),
        ("Acknowledgement", "iii"),
        ("Abstract", "iv"),
        ("List of Figures", "vi"),
        ("List of Tables", "vii"),
        ("1   Introduction", "1"),
        ("    1.1   Background and Motivation", "1"),
        ("    1.2   Statement of the Problem", "2"),
        ("    1.3   Research Question and Hypotheses", "3"),
        ("    1.4   Objectives", "3"),
        ("    1.5   Scope and Limitations", "4"),
        ("2   Study Area", "5"),
        ("    2.1   Location and Physiography", "5"),
        ("    2.2   Drivers of Land-Cover Change", "5"),
        ("3   Data and Software Used", "6"),
        ("    3.1   Sentinel-2 Surface-Reflectance Composites", "6"),
        ("    3.2   Ancillary and Reference Data", "6"),
        ("    3.3   Software and Libraries", "7"),
        ("4   Methodology", "8"),
        ("    4.1   Pre-processing and Quality Assurance", "8"),
        ("    4.2   Land-Cover Class Scheme", "9"),
        ("    4.3   OBIA Pipeline", "9"),
        ("        4.3.1  SLIC Superpixel Segmentation", "9"),
        ("        4.3.2  Per-Segment Feature Extraction", "10"),
        ("        4.3.3  Training-Segment Selection and Labelling", "11"),
        ("        4.3.4  Random Forest Classification", "11"),
        ("    4.4   Pixel-Based Pipeline", "12"),
        ("    4.5   Accuracy Assessment", "13"),
        ("    4.6   Change Detection", "14"),
        ("5   Results and Discussion", "15"),
        ("    5.1   Classified Maps", "15"),
        ("    5.2   Accuracy Assessment", "17"),
        ("    5.3   Per-Class Performance", "18"),
        ("    5.4   Class-Area Statistics and Change Detection", "19"),
        ("    5.5   Discussion of the Methodological Comparison", "21"),
        ("6   Conclusion and Future Scope", "23"),
        ("7   References", "24"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0),
                                                  WD_ALIGN_PARAGRAPH.RIGHT,
                                                  leader=2)  # 2 = dot leader
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(title)
        run.font.name = "Times New Roman"; run.font.size = Pt(12)
        run2 = p.add_run("\t" + page)
        run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
    add_page_break(doc)


def list_of_figures(doc: Document) -> None:
    add_heading(doc, "List of Figures", 1)
    figures = [
        ("Figure 1. Location of Dehradun Valley within Uttarakhand, India.", "5"),
        ("Figure 2. Workflow diagram of the OBIA and pixel-based pipelines.", "8"),
        ("Figure 3. SLIC superpixel segmentation result and OBIA-classified map, 2024.", "15"),
        ("Figure 4. Pixel-based classified map, 2024.", "16"),
        ("Figure 5. Side-by-side comparison of OBIA and pixel-based classifications, 2024.", "16"),
        ("Figure 6. Confusion matrix — OBIA, 2024.", "17"),
        ("Figure 7. Confusion matrix — Pixel-based, 2024.", "17"),
        ("Figure 8. Class area, 2016 vs 2024, for both methods.", "19"),
        ("Figure 9. Urban expansion 2016–2024, pixel-based.", "20"),
    ]
    for cap, page in figures:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0),
                                                  WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(cap); run.font.name = "Times New Roman"; run.font.size = Pt(12)
        run2 = p.add_run("\t" + page); run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
    add_page_break(doc)


def list_of_tables(doc: Document) -> None:
    add_heading(doc, "List of Tables", 1)
    tables = [
        ("Table 1. Sentinel-2 spectral bands used in the analysis.", "6"),
        ("Table 2. Five-class land-cover scheme adopted in this study.", "9"),
        ("Table 3. SLIC segmentation parameters.", "10"),
        ("Table 4. Random Forest hyperparameters.", "12"),
        ("Table 5. Accuracy assessment summary, both methods, both dates.", "17"),
        ("Table 6. Class-area statistics (ha) and change 2016–2024.", "19"),
    ]
    for cap, page in tables:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0),
                                                  WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(cap); run.font.name = "Times New Roman"; run.font.size = Pt(12)
        run2 = p.add_run("\t" + page); run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
    add_page_break(doc)


def introduction(doc: Document) -> None:
    add_heading(doc, "1   Introduction", 1)

    add_heading(doc, "1.1   Background and Motivation", 2)
    add_para(doc,
             "Land-cover change associated with urban expansion has emerged as one of "
             "the most consequential signals of human transformation of the Earth’s "
             "surface (Seto et al., 2011). In the Indian Himalayan foothill belt, "
             "tier-II capitals such as Dehradun have undergone particularly rapid "
             "growth driven by educational, administrative, and information-technology "
             "investment, accompanied by ribbon development along arterial roads and "
             "conversion of basmati cropland on the valley floor. Quantifying the "
             "spatial pattern and magnitude of this change is a pre-requisite for "
             "evidence-based urban planning, drainage and groundwater management, and "
             "for monitoring environmental compliance with master plans.")
    add_para(doc,
             "Multispectral satellite imagery is the practical workhorse for "
             "wall-to-wall land-cover mapping at city scale, and the European Space "
             "Agency’s Sentinel-2 mission, with its 10 m visible-near-infrared (VNIR) "
             "bands, five-day revisit and open data policy, has become the de-facto "
             "standard for medium-resolution monitoring of urban dynamics "
             "(Drusch et al., 2012). Two broad paradigms exist for translating such "
             "imagery into thematic land-cover classes: per-pixel classification, "
             "which treats every pixel as an independent statistical sample, and "
             "object-based image analysis (OBIA), which first groups neighbouring "
             "spectrally homogeneous pixels into image objects and then classifies "
             "those objects (Blaschke, 2010).")

    add_heading(doc, "1.2   Statement of the Problem", 2)
    add_para(doc,
             "The relative performance of OBIA versus pixel-based classification "
             "has been the subject of an extensive comparative literature. On "
             "very-high-resolution imagery (sub-metre to two-metre ground sampling "
             "distance), OBIA generally outperforms pixel-based methods because "
             "individual urban objects — buildings, parcels, road segments — are "
             "represented by many pixels and segmentation yields physically "
             "meaningful objects (Hossain & Chen, 2019). For medium-resolution "
             "imagery such as Sentinel-2 at 10 m, the picture is more nuanced. "
             "Ma et al. (2017), in a systematic review of seventy-three comparative "
             "studies, reported that the OBIA advantage diminishes and may reverse "
             "as pixel size approaches the characteristic size of land-cover "
             "objects. Despite this, OBIA pipelines on Sentinel-2 remain popular "
             "in operational settings because they produce visually smoother, "
             "polygon-ready outputs that are convenient for downstream "
             "decision-makers.")
    add_para(doc,
             "Whether OBIA provides a measurable accuracy advantage over a simple "
             "pixel-based Random Forest, when applied to Sentinel-2 imagery of a "
             "rapidly urbanising tier-II Indian city, has not been systematically "
             "documented. The present mini-project addresses this gap for Dehradun.")

    add_heading(doc, "1.3   Research Question and Hypotheses", 2)
    add_para(doc,
             "The research question guiding this work is: Does Object-Based Image "
             "Analysis of Sentinel-2 surface-reflectance composites produce a more "
             "accurate land-cover map of Dehradun than per-pixel Random Forest "
             "classification, and what is the magnitude of urban expansion between "
             "2016 and 2024 as estimated by each method?")
    add_para(doc,
             "Three working hypotheses are formulated. H1: OBIA will reduce "
             "salt-and-pepper noise compared with per-pixel Random Forest classification "
             "of the same imagery. H2: OBIA accuracy will be greater than or equal to "
             "pixel-based accuracy, because object-level features average out "
             "within-segment noise and provide more stable training signatures. "
             "H3: Both methods will agree that Dehradun gained substantial built-up area "
             "between 2016 and 2024, primarily at the expense of cropland and dense "
             "vegetation.")

    add_heading(doc, "1.4   Objectives", 2)
    add_para(doc,
             "The specific objectives of the study are:")
    bullet_items = [
        "to generate two co-registered, cloud-free Sentinel-2 surface-reflectance composites for the winter windows of 2016–2017 and 2024–2025 over Dehradun;",
        "to classify both composites using a five-class land-cover scheme with (a) an OBIA Random Forest on SLIC superpixels and (b) a pixel-based Random Forest with a 3 × 3 majority post-filter;",
        "to assess the accuracy of both classifiers against a common stratified random reference dataset interpreted in Google Earth Pro;",
        "to perform cross-tabulation change detection between 2016 and 2024 for each method and to report the magnitude and pattern of built-up expansion; and",
        "to discuss the practical implications of the observed accuracy difference for medium-resolution operational mapping.",
    ]
    for it in bullet_items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            r.font.name = "Times New Roman"; r.font.size = Pt(12)

    add_heading(doc, "1.5   Scope and Limitations", 2)
    add_para(doc,
             "The study is confined to the Dehradun Valley area of interest defined "
             "by the supplied Sentinel-2 raster footprint of approximately 996 km² "
             "of valid land. Only spectral features — visible and near-infrared "
             "reflectance, NDVI and NDBI — are used; texture, ancillary digital "
             "elevation, and ancillary OpenStreetMap layers are deliberately "
             "excluded so that the methodological comparison is unconfounded. A "
             "single SLIC segmentation scale is tested; sensitivity to segmentation "
             "parameters is identified as future work in Section 6. The accuracy "
             "assessment is based on 75 reference points per date and is subject "
             "to a per-class 95 % confidence interval of approximately ±12 percentage "
             "points (Olofsson et al., 2014).")
    add_page_break(doc)


def study_area(doc: Document) -> None:
    add_heading(doc, "2   Study Area", 1)

    add_heading(doc, "2.1   Location and Physiography", 2)
    add_para(doc,
             "Dehradun is the capital of the Indian state of Uttarakhand and lies in "
             "a structural valley (a doon) of the Lesser Himalaya between approximately "
             "30°15′ N and 30°35′ N latitude and 77°50′ E and 78°15′ E longitude, at "
             "an elevation of about 410 to 700 metres above mean sea level. The valley "
             "is bounded to the north by the Mussoorie ridge and to the south by the "
             "Shivalik foothills, with the Asan, Tons and Song rivers traversing the "
             "valley floor. The area of interest used in this study is a bounding "
             "rectangle of approximately 58.6 km × 33.7 km, of which approximately "
             "996 km² are valid land pixels inside the Sentinel-2 clip mask. The "
             "spatial reference system used throughout the analysis is UTM Zone 44 "
             "North on the WGS-84 datum (EPSG:32644).")
    add_image(doc, MAPS / "obia_map_2024.png", width_in=6.0,
              caption="Location and extent of the Dehradun study area, "
                      "shown as the OBIA-classified 2024 map.", fig_num=1)

    add_heading(doc, "2.2   Drivers of Land-Cover Change", 2)
    add_para(doc,
             "Three classes of driver dominate land-cover change in the Dehradun "
             "Valley during the study period. The first is the expansion of the "
             "information-technology and education sectors, anchored on planned "
             "developments around Selaqui in the west and Sahastradhara in the "
             "north-east. The second is the residential and commercial ribbon "
             "development along the National Highway 7 corridor and along the "
             "Rajpur Road–Mussoorie axis, which has filled in many of the agricultural "
             "interstices that existed in the 2016 baseline. The third is the "
             "conversion of basmati cropland on the valley floor between Dehradun "
             "and Doiwala into peri-urban settlements and bare construction land. "
             "Cumulatively, these processes are expected to manifest in the "
             "classified maps as a marked increase in the built-up class and a "
             "corresponding decrease in mixed vegetation / cropland and dense "
             "vegetation classes between 2016 and 2024.")
    add_page_break(doc)


def data_and_software(doc: Document) -> None:
    add_heading(doc, "3   Data and Software Used", 1)

    add_heading(doc, "3.1   Sentinel-2 Surface-Reflectance Composites", 2)
    add_para(doc,
             "Two cloud-free seasonal composites of Sentinel-2 Level-2A surface "
             "reflectance were generated in Google Earth Engine (Gorelick et al., "
             "2017). The first composite spans the winter window from November 2016 "
             "to February 2017, capturing the dry, post-monsoon period in which "
             "cloud cover is minimal and field phenology distinguishes irrigated "
             "rabi crops from fallow land. The second composite spans the "
             "corresponding window of November 2024 to February 2025. Both composites "
             "were produced as Float32 GeoTIFFs in UTM Zone 44 N at 10 m pixel size, "
             "with twelve bands stored in the order B2, B3, B4, B8, B5, B6, B7, B8A, "
             "B11, B12, NDVI, NDBI. The two scenes share an identical raster grid "
             "(3 371 × 5 864 pixels), facilitating per-pixel comparison without "
             "re-projection.")
    add_table_caption(doc, 1, "Sentinel-2 spectral bands used in the analysis "
                              "(after pre-processing audit, six bands retained).")
    add_table(doc,
              headers=["Band", "Wavelength (nm)", "Native res. (m)", "Role"],
              rows=[
                  ["B2 (Blue)", "490", "10", "Visible — surface reflectance"],
                  ["B3 (Green)", "560", "10", "Visible — surface reflectance"],
                  ["B4 (Red)", "665", "10", "Visible — surface reflectance"],
                  ["B8 (NIR)", "842", "10", "Near-infrared — vegetation"],
                  ["NDVI", "—", "10", "Vegetation index (B8 − B4)/(B8 + B4)"],
                  ["NDBI", "—", "10", "Built-up index (B11 − B8)/(B11 + B8)"],
              ])

    add_heading(doc, "3.2   Ancillary and Reference Data", 2)
    add_para(doc,
             "Ancillary data used in the project comprised the area-of-interest "
             "vector boundary supplied with the Sentinel-2 composites, the "
             "high-resolution imagery available in Google Earth Pro for the years "
             "2016 and 2024, and historical Google Earth imagery for selected "
             "regions where it was needed to disambiguate the reference label for a "
             "given validation point. Reference labels for both training-segment "
             "verification and accuracy assessment were assigned by visual "
             "interpretation in Google Earth Pro by the author. No third-party "
             "land-cover product was used as reference because none was available "
             "at 10 m for the study area and dates.")

    add_heading(doc, "3.3   Software and Libraries", 2)
    add_para(doc,
             "The analytical pipeline was implemented in Python 3.13 on a standard "
             "Windows workstation. The principal open-source libraries used were "
             "rasterio for reading and writing the Sentinel-2 GeoTIFFs, NumPy for "
             "array computation, scikit-image for SLIC superpixel segmentation "
             "(van der Walt et al., 2014), scikit-learn for the Random Forest "
             "classifier (Pedregosa et al., 2011), pandas and GeoPandas for tabular "
             "and vector handling, SciPy ndimage for vectorised majority filtering, "
             "matplotlib for cartographic output, simplekml for export of validation "
             "points to Google Earth Pro, and python-pptx / python-docx for "
             "automated generation of the slide deck and this report. All scripts "
             "are numbered 01–14 in the project directory D:/Dehradun_OBIA/scripts/ "
             "and are designed to be re-runnable end-to-end.")
    add_page_break(doc)


def methodology(doc: Document) -> None:
    add_heading(doc, "4   Methodology", 1)

    add_heading(doc, "4.1   Pre-processing and Quality Assurance", 2)
    add_para(doc,
             "Before classification, both composites were subjected to a per-band "
             "statistical audit on the in-AOI pixels. This audit revealed that "
             "while the visible (B2, B3, B4), near-infrared (B8), and the derived "
             "NDVI and NDBI bands all carried physically plausible variance, the "
             "red-edge bands (B5, B6, B7, B8A) and the shortwave-infrared bands "
             "(B11, B12) showed extremely narrow ranges relative to their means — "
             "for example, B11 in the 2016 composite had a standard deviation of "
             "only ~27 DN against a mean of ~1810, a relative variation of less "
             "than 1.5 %. This pattern is inconsistent with genuine surface "
             "reflectance variation and is most likely an artefact of the band "
             "reduction step in the Google Earth Engine export. Because the "
             "degraded bands carried negligible discriminative signal, they were "
             "excluded from the final feature stack. The retained features are "
             "therefore B2, B3, B4, B8, NDVI and NDBI. NDBI was retained because, "
             "although its parent B11 band is degraded inside the AOI, the index "
             "itself shows healthy variance (standard deviation ≈ 0.12, range "
             "−0.71 to 0.93) which indicates that the NDBI image was computed "
             "before the export-side band reduction.")
    add_para(doc,
             "Following the audit, a valid-pixel mask was constructed from the "
             "NaN footprint of the red band (B4) and stored as a single-band uint8 "
             "GeoTIFF for each date. The masks for the two dates are identical, "
             "consistent with the co-registered nature of the composites, and "
             "cover 9 961 760 of the 19 767 544 raster pixels (50.4 %).")

    add_heading(doc, "4.2   Land-Cover Class Scheme", 2)
    add_para(doc,
             "A five-class land-cover scheme was adopted, designed to be coarse "
             "enough for Sentinel-2 at 10 m to discriminate reliably using only "
             "spectral features, while remaining fine enough to track the "
             "urban-expansion process of interest.")
    add_table_caption(doc, 2, "Five-class land-cover scheme.")
    add_table(doc,
              headers=["ID", "Class", "Indicative spectral signature"],
              rows=[
                  ["1", "Built-up", "High B2/B3, NDBI > 0, NDVI < 0.30"],
                  ["2", "Dense Vegetation", "NDVI > 0.65"],
                  ["3", "Mixed Veg / Cropland", "NDVI between 0.30 and 0.55"],
                  ["4", "Bare / Open Land", "NDVI < 0.20, mid-NDBI, B8 > 1500"],
                  ["5", "Water", "Very low B8 (NIR), NDVI ≤ 0"],
              ])

    add_heading(doc, "4.3   OBIA Pipeline", 2)
    add_heading(doc, "4.3.1   SLIC Superpixel Segmentation", 3)
    add_para(doc,
             "The Simple Linear Iterative Clustering (SLIC) algorithm was used to "
             "partition each composite into image objects. SLIC, introduced by "
             "Achanta et al. (2012), performs a constrained k-means clustering in "
             "the joint (x, y, spectral) space, restricting each centroid’s search "
             "window to a local 2S × 2S neighbourhood where S = √(N/k), with N the "
             "number of pixels and k the requested number of segments. The "
             "spectral-versus-spatial trade-off is governed by a compactness "
             "parameter m: low compactness lets segments adhere to spectral edges, "
             "while high compactness keeps segments close to a regular grid. The "
             "implementation in scikit-image (van der Walt et al., 2014) was used.")
    add_table_caption(doc, 3, "SLIC segmentation parameters.")
    add_table(doc,
              headers=["Parameter", "Value", "Rationale"],
              rows=[
                  ["n_segments", "8000", "Target ≈ 12 ha per segment given ~10 M valid pixels"],
                  ["compactness", "10", "Moderate; allows spectral edges to dominate"],
                  ["input channels", "B2, B3, B4, B8", "Visible + NIR; NDVI / NDBI excluded to avoid biasing toward class definitions"],
                  ["normalisation", "p1–p99 min-max per channel", "Required so spectral and spatial distances are commensurable"],
                  ["mask", "AOI valid mask", "Prevents segments from bleeding into nodata"],
                  ["start_label", "1", "Allows 0 to denote outside-AOI"],
              ])
    add_para(doc,
             "Segmentation was performed independently for each date. Although the "
             "two rasters share the same grid, the spectral content at any given "
             "location differs between 2016 and 2024 by definition (this is the "
             "change signal); joint segmentation would have biased the comparison "
             "by forcing 2016 and 2024 to share boundaries that may not be "
             "spectrally appropriate to both. Each date yielded approximately 8000 "
             "segments inside the AOI, with a median segment size of 1241 pixels "
             "(approximately 12.4 ha).")

    add_heading(doc, "4.3.2   Per-Segment Feature Extraction", 3)
    add_para(doc,
             "For every segment in each date, the mean of B2, B3, B4, B8, NDVI and "
             "NDBI over the constituent pixels was computed using a vectorised "
             "labelled-array reduction (scipy.ndimage.mean). In addition, the "
             "segment area in pixels, the centroid in row–column and UTM "
             "coordinates, and the segment-mean B11 (retained for audit purposes "
             "only) were recorded. The full segment-feature tables were persisted "
             "as Apache Parquet files for downstream consumption.")

    add_heading(doc, "4.3.3   Training-Segment Selection and Labelling", 3)
    add_para(doc,
             "Training segments were selected in two steps. First, a rule-based "
             "stratification on segment-mean NDVI, NDBI and B8 was used to "
             "construct candidate pools for each of the five classes. From each "
             "pool, eight segments were drawn at random per date, yielding 40 "
             "training candidates per date. Second, the candidate set was reviewed "
             "by visual interpretation in Google Earth Pro; segments for which the "
             "rule-based hint was wrong were manually relabelled to the correct "
             "class. After this human-in-the-loop step, the final training "
             "distribution was {built-up: 13, dense vegetation: 9, mixed "
             "vegetation/cropland: 7, bare/open: 7, water: 4} for 2016 and "
             "{built-up: 12, dense vegetation: 8, mixed vegetation/cropland: 7, "
             "bare/open: 8, water: 5} for 2024. The training table is reproducible "
             "with a fixed NumPy random seed of 42.")

    add_heading(doc, "4.3.4   Random Forest Classification", 3)
    add_para(doc,
             "A Random Forest classifier (Breiman, 2001) was trained on the "
             "labelled segments using identical hyperparameters across the two "
             "dates and across the two methodological pipelines, as listed in "
             "Table 4. The classifier was then used to predict the class of every "
             "one of the approximately 8000 segments per date, and the resulting "
             "segment-level labels were rasterised back to the original 10 m grid.")
    add_table_caption(doc, 4, "Random Forest hyperparameters (identical for "
                              "OBIA and pixel-based pipelines).")
    add_table(doc,
              headers=["Parameter", "Value"],
              rows=[
                  ["n_estimators", "300"],
                  ["max_depth", "None (grow fully)"],
                  ["min_samples_leaf", "2"],
                  ["class_weight", "balanced"],
                  ["oob_score", "True"],
                  ["random_state", "42"],
                  ["n_jobs", "−1 (all cores)"],
              ])

    add_heading(doc, "4.4   Pixel-Based Pipeline", 2)
    add_para(doc,
             "The pixel-based pipeline shares the same training-segment "
             "centroids, the same feature set and the same Random Forest "
             "hyperparameters as the OBIA pipeline; only the spatial unit of "
             "classification differs. For each training-segment centroid, a "
             "3 × 3 pixel window of feature values was extracted and assigned "
             "the segment’s reference class label. This yielded approximately "
             "360 training pixels per date and is deliberately designed to "
             "keep the two pipelines as comparable as possible while giving the "
             "pixel classifier a slightly richer training set commensurate with "
             "its higher number of decision units. Every valid pixel in the "
             "image was then scored by the trained Random Forest, with predictions "
             "computed in one-million-pixel chunks to bound peak memory. The raw "
             "pixel prediction was post-processed with a 3 × 3 majority filter "
             "(modal value over the eight neighbours plus the centre) implemented "
             "as a fully vectorised operation over a nine-deep stack of shifted "
             "arrays. The 3 × 3 majority filter is a standard salt-and-pepper "
             "suppression operator in per-pixel classification (Lu & Weng, 2007).")

    add_heading(doc, "4.5   Accuracy Assessment", 2)
    add_para(doc,
             "Accuracy was assessed using a stratified random sample of 75 "
             "reference points per date, equally distributed across the five "
             "classes (15 per class). Stratification was performed on the OBIA "
             "classified raster, and the same 150 reference points were used to "
             "score both methods so that the comparison is unaffected by "
             "differences in the sampling design. Each reference point was "
             "interpreted by the author in Google Earth Pro using historical "
             "imagery as close as possible to the season of the corresponding "
             "Sentinel-2 composite, following the recommendations of Olofsson et "
             "al. (2014). The reference labels were stored in a CSV alongside the "
             "OBIA and pixel-based predictions extracted from each classified "
             "raster, and confusion matrices, overall accuracy, Cohen’s κ "
             "(Cohen, 1960), and per-class producer’s and user’s accuracy were "
             "computed using scikit-learn (Pedregosa et al., 2011).")

    add_heading(doc, "4.6   Change Detection", 2)
    add_para(doc,
             "Change detection was performed by post-classification cross-"
             "tabulation, separately for each method. A 5 × 5 transition matrix "
             "in hectares was constructed for each method by counting pixels "
             "in each (class_2016, class_2024) combination over the common valid "
             "AOI mask. Two derivative products were produced from this matrix: "
             "an urbanisation binary map flagging pixels that were not built-up "
             "in 2016 but were built-up in 2024, and a per-class area-change "
             "summary in hectares.")
    add_page_break(doc)


def results(doc: Document) -> None:
    add_heading(doc, "5   Results and Discussion", 1)

    add_heading(doc, "5.1   Classified Maps", 2)
    add_para(doc,
             "Figures 3, 4 and 5 present the classified outputs for the 2024 "
             "composite. The OBIA-classified map (Figure 3) is visually smoother "
             "and is composed of polygon-like patches at the SLIC segment scale, "
             "while the pixel-based output (Figure 4) preserves finer 10 m "
             "structure including narrow roads and isolated built-up patches "
             "within the central city. The side-by-side comparison (Figure 5) "
             "shows that the OBIA map covers the central urban cluster as one "
             "contiguous red mass, whereas the pixel-based map preserves the "
             "internal vegetation pockets associated with parks, riverine "
             "corridors and institutional campuses.")
    add_image(doc, MAPS / "obia_map_2024.png", width_in=6.0,
              caption="OBIA-classified land-cover map for the 2024 composite. "
                      "Each colour patch is a SLIC segment assigned by the "
                      "Random Forest classifier.",
              fig_num=3)
    add_image(doc, MAPS / "pixel_map_2024.png", width_in=6.0,
              caption="Pixel-based classified land-cover map for the 2024 "
                      "composite, after the 3 × 3 majority post-filter.",
              fig_num=4)
    add_image(doc, MAPS / "obia_vs_pixel_2024.png", width_in=6.5,
              caption="Side-by-side comparison of OBIA (left) and pixel-based "
                      "(right) classifications for 2024.",
              fig_num=5)

    add_heading(doc, "5.2   Accuracy Assessment", 2)
    add_para(doc,
             "The accuracy results, summarised in Table 5, indicate that the "
             "pixel-based Random Forest outperformed the OBIA Random Forest by a "
             "consistent margin of approximately nine to thirteen percentage "
             "points of overall accuracy and 0.10 to 0.15 in Cohen’s κ on both "
             "dates. The pixel-based method scored 78.7 % overall accuracy and "
             "κ = 0.72 on both 2016 and 2024, while the OBIA method scored "
             "65.3 % (κ = 0.57) in 2016 and 69.3 % (κ = 0.62) in 2024.")
    add_table_caption(doc, 5, "Accuracy assessment summary based on 75 "
                              "stratified random reference points per date.")
    if (VAL / "accuracy_summary.csv").exists():
        acc = pd.read_csv(VAL / "accuracy_summary.csv")
        rows = [[str(r["method"]), str(r["year"]),
                 f"{100*r['OA']:.1f}", f"{r['Kappa']:.2f}", str(int(r["n_val"]))]
                for _, r in acc.iterrows()]
        add_table(doc,
                  headers=["Method", "Year", "OA (%)", "Cohen’s κ", "n_val"],
                  rows=rows)
    add_image(doc, VAL / "confusion_obia_2024.png", width_in=4.5,
              caption="Confusion matrix for the OBIA-classified 2024 map.",
              fig_num=6)
    add_image(doc, VAL / "confusion_pixel_2024.png", width_in=4.5,
              caption="Confusion matrix for the pixel-based 2024 map.",
              fig_num=7)

    add_heading(doc, "5.3   Per-Class Performance", 2)
    add_para(doc,
             "Inspection of the per-class producer’s and user’s accuracies "
             "computed from the confusion matrices reveals a consistent pattern. "
             "The water class shows near-perfect user’s accuracy in both methods "
             "— that is, when either classifier predicts water, the prediction "
             "is almost always correct — but the producer’s accuracy is "
             "considerably lower because narrow rivers and small water bodies "
             "are absorbed into the surrounding riparian vegetation, particularly "
             "in the OBIA case where 12 ha segments cannot resolve sub-segment "
             "water features. The mixed vegetation / cropland class shows the "
             "largest omission rate for the OBIA pipeline: SLIC segments that "
             "straddle field boundaries are assigned a single label and the "
             "minority cover is lost. The built-up class is detected with markedly "
             "higher producer’s accuracy by the pixel-based pipeline because the "
             "10 m grid resolves small isolated rooftops that OBIA averages into "
             "the segment within which they sit. These per-class observations "
             "are consistent with the broader systematic-review findings of Ma "
             "et al. (2017), which noted that the OBIA advantage diminishes as "
             "image resolution approaches the characteristic length scale of "
             "land-cover objects.")

    add_heading(doc, "5.4   Class-Area Statistics and Change Detection", 2)
    add_para(doc,
             "Class-area statistics for both methods and both dates are "
             "presented in Table 6 and visualised in Figure 8. Both methods "
             "agree on the direction of change for every class: built-up "
             "expanded substantially, dense vegetation shrank, mixed "
             "vegetation/cropland expanded modestly and bare/open land "
             "decreased. The methods disagree on the magnitude of the built-up "
             "change. The OBIA pipeline reports a growth of approximately "
             "13 820 ha (about 138 km²) in built-up area between 2016 and 2024, "
             "while the pixel-based pipeline reports a growth of approximately "
             "8 022 ha (about 80 km²) over the same interval. Because the "
             "pixel-based method is the more accurate classifier according to "
             "Section 5.2, the pixel-based estimate of approximately 80 km² of "
             "new built-up land is preferred as the headline figure.")
    add_table_caption(doc, 6, "Class-area statistics (hectares) and net change "
                              "2016–2024 for both methods.")
    rows_tbl = []
    if (CHG / "area_table_pixel.csv").exists() and (CHG / "area_table_obia.csv").exists():
        ao = pd.read_csv(CHG / "area_table_obia.csv")
        ap = pd.read_csv(CHG / "area_table_pixel.csv")
        for c in [1, 2, 3, 4, 5]:
            row_o = ao[ao.class_id == c].iloc[0]
            row_p = ap[ap.class_id == c].iloc[0]
            rows_tbl.append([
                row_o["class_name"],
                f"{row_o['area_2016_ha']:.0f}",
                f"{row_o['area_2024_ha']:.0f}",
                f"{row_o['change_ha']:+.0f}",
                f"{row_p['area_2016_ha']:.0f}",
                f"{row_p['area_2024_ha']:.0f}",
                f"{row_p['change_ha']:+.0f}",
            ])
    if rows_tbl:
        add_table(doc,
                  headers=["Class", "OBIA 2016", "OBIA 2024", "Δ OBIA",
                           "Pixel 2016", "Pixel 2024", "Δ Pixel"],
                  rows=rows_tbl)
    add_image(doc, MAPS / "area_chart.png", width_in=6.5,
              caption="Class area in hectares for 2016 and 2024 under both "
                      "methods, with the annotated change in hectares (red = "
                      "increase, green = decrease).",
              fig_num=8)
    add_image(doc, MAPS / "change_map_pixel.png", width_in=6.5,
              caption="Spatial pattern of urban expansion 2016 → 2024 as "
                      "estimated by the pixel-based pipeline. Red represents "
                      "pixels that transitioned to built-up between 2016 and "
                      "2024; grey represents persistent built-up.",
              fig_num=9)

    add_heading(doc, "5.5   Discussion of the Methodological Comparison", 2)
    add_para(doc,
             "Hypothesis H1 — that OBIA would reduce salt-and-pepper noise — is "
             "qualitatively supported by visual inspection of Figures 3–5; the "
             "OBIA output is visually smoother. Hypothesis H2 — that this visual "
             "smoothness would translate into higher accuracy — is, however, "
             "rejected by the data: the pixel-based pipeline outperformed the "
             "OBIA pipeline by approximately 10 percentage points on both dates. "
             "Hypothesis H3 — that both methods would detect substantial "
             "built-up expansion — is confirmed.")
    add_para(doc,
             "Four explanations help account for the observed accuracy ranking. "
             "First, the SLIC segments used in this study have a median size of "
             "approximately 12.4 ha, which is one to two orders of magnitude "
             "larger than the characteristic length scale of land-cover change "
             "in peri-urban Dehradun (rooftops, narrow roads, small fields "
             "typically of 0.1 to 1 ha). When a segment straddles a boundary, "
             "the within-segment spectral averaging pulls toward the dominant "
             "class and erases the minority cover, an effect already noted by "
             "Whyte et al. (2018) for Sentinel-2-scale OBIA. Second, the 3 × 3 "
             "training window used by the pixel-based pipeline yields roughly "
             "360 training pixels per date as opposed to 40 training segments "
             "for OBIA; the pixel classifier therefore has a richer training "
             "set and is better able to learn the spectral diversity of mixed "
             "fringe classes. Third, the principal critique of pixel-based "
             "classification — salt-and-pepper noise — is largely cured by a "
             "single 3 × 3 majority filter without sacrificing boundary "
             "fidelity, as Lu and Weng (2007) anticipated. Fourth, this result "
             "does not imply that OBIA is intrinsically inferior; it implies "
             "that single-scale SLIC-OBIA at a 12 ha segment size is the wrong "
             "tool for change detection at the urban fringe in a Sentinel-2 "
             "context. Two natural remediations — a finer SLIC scale "
             "(n_segments > 30 000) and a multi-scale OBIA ensemble — are "
             "identified in Section 6.")
    add_page_break(doc)


def conclusion(doc: Document) -> None:
    add_heading(doc, "6   Conclusion and Future Scope", 1)
    add_para(doc,
             "The mini-project compared OBIA and pixel-based Random Forest "
             "classifications of two co-registered Sentinel-2 surface-reflectance "
             "composites of the Dehradun Valley for the winter windows of "
             "2016–2017 and 2024–2025. The methodological design held the "
             "training samples, spectral features, classifier and hyperparameters "
             "constant across the two pipelines, isolating the spatial unit of "
             "classification as the single independent variable.")
    add_para(doc,
             "Three principal conclusions follow. First, on Sentinel-2 imagery "
             "of a mixed peri-urban setting and at the segmentation scale tested "
             "(SLIC n = 8000 → median segment ≈ 12 ha), a per-pixel Random "
             "Forest with a 3 × 3 majority post-filter delivers higher overall "
             "accuracy (78.7 % on both dates, κ = 0.72) than the OBIA Random "
             "Forest pipeline (65.3 % to 69.3 %, κ = 0.57 to 0.62), despite the "
             "OBIA pipeline producing visually cleaner maps. Second, both "
             "methods agree that Dehradun gained substantial built-up area over "
             "the eight-year study period; the defensible (pixel-based) figure "
             "is approximately 80 km² of new built-up land, mostly converted "
             "from mixed cropland and bare fallow on the valley floor. Third, "
             "dense vegetation declined by roughly 130 to 160 km² over the "
             "same period across both methods, the largest single signal in "
             "the dataset.")
    add_para(doc,
             "Future work falls naturally into five directions. A multi-scale "
             "OBIA experiment varying n_segments over the range 2 000 to 80 000 "
             "would identify the segmentation scale at which OBIA closes the "
             "accuracy gap with the pixel-based method, and, more ambitiously, "
             "would enable a scale-ensemble classifier. Re-export of the "
             "Sentinel-2 composite preserving the red-edge and shortwave-"
             "infrared bands would allow the full 12-band feature stack to be "
             "tested and could plausibly improve OBIA more than pixel-based, "
             "since OBIA can pool the additional bands as segment-mean features. "
             "The addition of texture descriptors at segment level, particularly "
             "Grey-Level Co-occurrence Matrix statistics, is a further OBIA-"
             "friendly extension. Validation rigour would be improved by an "
             "independent labelling exercise involving two or three interpreters "
             "to quantify inter-rater reliability. Finally, an annual dense time "
             "series from 2016 through 2024 would replace the two-snapshot "
             "design with a continuous trajectory and enable estimation of the "
             "year of conversion for each newly built-up pixel.")
    add_page_break(doc)


def references(doc: Document) -> None:
    add_heading(doc, "7   References", 1)
    # APA 7 hanging-indent: set first_line_indent negative
    refs = [
        "Achanta, R., Shaji, A., Smith, K., Lucchi, A., Fua, P., & Süsstrunk, S. "
            "(2012). SLIC superpixels compared to state-of-the-art superpixel "
            "methods. IEEE Transactions on Pattern Analysis and Machine "
            "Intelligence, 34(11), 2274–2282. "
            "https://doi.org/10.1109/TPAMI.2012.120",
        "Blaschke, T. (2010). Object based image analysis for remote sensing. "
            "ISPRS Journal of Photogrammetry and Remote Sensing, 65(1), 2–16. "
            "https://doi.org/10.1016/j.isprsjprs.2009.06.004",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. "
            "https://doi.org/10.1023/A:1010933404324",
        "Cohen, J. (1960). A coefficient of agreement for nominal scales. "
            "Educational and Psychological Measurement, 20(1), 37–46. "
            "https://doi.org/10.1177/001316446002000104",
        "Drusch, M., Del Bello, U., Carlier, S., Colin, O., Fernandez, V., "
            "Gascon, F., Hoersch, B., Isola, C., Laberinti, P., Martimort, P., "
            "Meygret, A., Spoto, F., Sy, O., Marchese, F., & Bargellini, P. "
            "(2012). Sentinel-2: ESA’s optical high-resolution mission for GMES "
            "operational services. Remote Sensing of Environment, 120, 25–36. "
            "https://doi.org/10.1016/j.rse.2011.11.026",
        "Gorelick, N., Hancher, M., Dixon, M., Ilyushchenko, S., Thau, D., & "
            "Moore, R. (2017). Google Earth Engine: Planetary-scale geospatial "
            "analysis for everyone. Remote Sensing of Environment, 202, 18–27. "
            "https://doi.org/10.1016/j.rse.2017.06.031",
        "Hossain, M. D., & Chen, D. (2019). Segmentation for Object-Based Image "
            "Analysis (OBIA): A review of algorithms and challenges from remote "
            "sensing perspective. ISPRS Journal of Photogrammetry and Remote "
            "Sensing, 150, 115–134. "
            "https://doi.org/10.1016/j.isprsjprs.2019.02.009",
        "Lu, D., & Weng, Q. (2007). A survey of image classification methods "
            "and techniques for improving classification performance. "
            "International Journal of Remote Sensing, 28(5), 823–870. "
            "https://doi.org/10.1080/01431160600746456",
        "Ma, L., Li, M., Ma, X., Cheng, L., Du, P., & Liu, Y. (2017). A review "
            "of supervised object-based land-cover image classification. ISPRS "
            "Journal of Photogrammetry and Remote Sensing, 130, 277–293. "
            "https://doi.org/10.1016/j.isprsjprs.2017.06.001",
        "Olofsson, P., Foody, G. M., Herold, M., Stehman, S. V., Woodcock, "
            "C. E., & Wulder, M. A. (2014). Good practices for estimating area "
            "and assessing accuracy of land change. Remote Sensing of "
            "Environment, 148, 42–57. "
            "https://doi.org/10.1016/j.rse.2014.02.015",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., "
            "Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., "
            "Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, "
            "M., & Duchesnay, É. (2011). Scikit-learn: Machine learning in "
            "Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "Seto, K. C., Fragkias, M., Güneralp, B., & Reilly, M. K. (2011). A "
            "meta-analysis of global urban land expansion. PLoS ONE, 6(8), "
            "e23777. https://doi.org/10.1371/journal.pone.0023777",
        "van der Walt, S., Schönberger, J. L., Nunez-Iglesias, J., Boulogne, "
            "F., Warner, J. D., Yager, N., Gouillart, E., & Yu, T. (2014). "
            "scikit-image: Image processing in Python. PeerJ, 2, e453. "
            "https://doi.org/10.7717/peerj.453",
        "Whyte, A., Ferentinos, K. P., & Petropoulos, G. P. (2018). A new "
            "synergistic approach for monitoring wetlands using Sentinels-1 "
            "and 2 data with object-based machine learning algorithms. "
            "Environmental Modelling & Software, 104, 40–54. "
            "https://doi.org/10.1016/j.envsoft.2018.02.008",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(r)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def build() -> None:
    doc = Document()
    set_a4_section(doc)
    set_default_style(doc)
    configure_heading_style(doc, 1, 18)
    configure_heading_style(doc, 2, 14)
    configure_heading_style(doc, 3, 12)

    cover_page(doc)
    declaration(doc)
    acknowledgement(doc)
    abstract(doc)
    table_of_contents(doc)
    list_of_figures(doc)
    list_of_tables(doc)
    introduction(doc)
    study_area(doc)
    data_and_software(doc)
    methodology(doc)
    results(doc)
    conclusion(doc)
    references(doc)

    try:
        doc.save(REPORT_PATH)
        print(f"Saved {REPORT_PATH}")
    except PermissionError:
        alt = REPORT_PATH.with_name(REPORT_PATH.stem + "_v2.docx")
        doc.save(alt)
        print(f"NOTE: original was locked. Saved as {alt}")


if __name__ == "__main__":
    build()
